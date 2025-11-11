#!/usr/bin/env python3
"""
Convert Markdown to DOCX format
Preserves headings, code blocks, tables, and basic formatting
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Installing required packages...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE

def setup_styles(doc):
    """Setup custom styles for the document"""
    styles = doc.styles

    # Code block style
    if 'Code Block' not in [s.name for s in styles]:
        code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(9)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)

def parse_markdown_line(line):
    """Parse a markdown line and return its type and content"""
    line = line.rstrip()

    # Headings
    if line.startswith('# '):
        return ('h1', line[2:])
    elif line.startswith('## '):
        return ('h2', line[3:])
    elif line.startswith('### '):
        return ('h3', line[4:])
    elif line.startswith('#### '):
        return ('h4', line[5:])
    elif line.startswith('##### '):
        return ('h5', line[6:])

    # Horizontal rule
    elif line.strip() in ['---', '***', '___']:
        return ('hr', '')

    # Bullet list
    elif line.startswith('- ') or line.startswith('* '):
        return ('bullet', line[2:])

    # Numbered list
    elif re.match(r'^\d+\.\s', line):
        match = re.match(r'^\d+\.\s(.+)', line)
        return ('number', match.group(1) if match else line)

    # Code block markers
    elif line.startswith('```'):
        return ('code_fence', line[3:])

    # Table separator
    elif re.match(r'^\|?[-:\s|]+\|?$', line.strip()):
        return ('table_sep', line)

    # Table row
    elif '|' in line:
        return ('table', line)

    # Empty line
    elif not line.strip():
        return ('empty', '')

    # Normal paragraph
    else:
        return ('para', line)

def apply_inline_formatting(paragraph, text):
    """Apply inline markdown formatting (bold, italic, code, links)"""
    # Process inline code first
    parts = re.split(r'(`[^`]+`)', text)

    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            # Inline code
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            # Process bold and italic
            segments = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__|_[^_]+_|\[([^\]]+)\]\(([^)]+)\))', part)

            i = 0
            while i < len(segments):
                seg = segments[i]

                if not seg:
                    i += 1
                    continue

                # Bold
                if (seg.startswith('**') and seg.endswith('**')) or (seg.startswith('__') and seg.endswith('__')):
                    run = paragraph.add_run(seg[2:-2])
                    run.bold = True
                # Italic
                elif (seg.startswith('*') and seg.endswith('*')) or (seg.startswith('_') and seg.endswith('_')):
                    run = paragraph.add_run(seg[1:-1])
                    run.italic = True
                # Link
                elif i + 2 < len(segments) and segments[i+1] and segments[i+2]:
                    link_text = segments[i+1]
                    link_url = segments[i+2]
                    run = paragraph.add_run(f"{link_text} ({link_url})")
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.underline = True
                    i += 2
                else:
                    paragraph.add_run(seg)

                i += 1

def convert_md_to_docx(md_file, docx_file):
    """Convert markdown file to docx"""
    print(f"Converting {md_file} to {docx_file}...")

    doc = Document()
    setup_styles(doc)

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    i = 0
    while i < len(lines):
        line = lines[i]
        line_type, content = parse_markdown_line(line)

        # Handle code blocks
        if line_type == 'code_fence':
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text, style='Code Block')
                code_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line.rstrip())
            i += 1
            continue

        # Handle tables
        if line_type == 'table' or line_type == 'table_sep':
            if line_type == 'table':
                cells = [cell.strip() for cell in line.split('|')]
                cells = [c for c in cells if c]  # Remove empty cells
                table_rows.append(cells)

            # Check if next line is still table
            if i + 1 < len(lines):
                next_type, _ = parse_markdown_line(lines[i + 1])
                if next_type not in ['table', 'table_sep']:
                    # End of table
                    if table_rows:
                        create_table(doc, table_rows)
                        table_rows = []
            elif table_rows:
                # Last line and we have table data
                create_table(doc, table_rows)
                table_rows = []

            i += 1
            continue

        # Handle other elements
        if line_type == 'h1':
            p = doc.add_heading(content, level=1)
        elif line_type == 'h2':
            p = doc.add_heading(content, level=2)
        elif line_type == 'h3':
            p = doc.add_heading(content, level=3)
        elif line_type == 'h4':
            p = doc.add_heading(content, level=4)
        elif line_type == 'h5':
            p = doc.add_heading(content, level=5)
        elif line_type == 'hr':
            p = doc.add_paragraph()
            p.add_run('_' * 80)
        elif line_type == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, content)
        elif line_type == 'number':
            p = doc.add_paragraph(style='List Number')
            apply_inline_formatting(p, content)
        elif line_type == 'empty':
            doc.add_paragraph()
        elif line_type == 'para':
            p = doc.add_paragraph()
            apply_inline_formatting(p, content)

        i += 1

    # Save document
    doc.save(docx_file)
    print(f"Successfully converted to {docx_file}")
    print(f"File size: {Path(docx_file).stat().st_size / 1024:.1f} KB")

def create_table(doc, rows):
    """Create a table from rows"""
    if not rows:
        return

    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Light Grid Accent 1'

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            if j < len(row.cells):
                row.cells[j].text = cell_data
                # Make header row bold
                if i == 0:
                    for paragraph in row.cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

if __name__ == '__main__':
    md_file = r'f:\FINALPROJECT\Codemaster-3\docs\academicreport\changes.md'
    docx_file = r'f:\FINALPROJECT\Codemaster-3\docs\academicreport\changes.docx'

    convert_md_to_docx(md_file, docx_file)
