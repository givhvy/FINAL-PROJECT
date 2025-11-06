#!/usr/bin/env python3
"""
Convert Markdown academic report to Word document with proper formatting
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def convert_md_to_docx(md_file, docx_file):
    """Convert markdown file to formatted Word document"""
    
    # Create new document
    doc = Document()
    
    # Set document margins (1 inch all sides for academic)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    print(f"Processing {len(lines)} lines from {md_file}...")
    
    in_code_block = False
    code_lines = []
    
    for i, line in enumerate(lines):
        line = line.rstrip('\n')
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block - add accumulated code
                if code_lines:
                    p = doc.add_paragraph('\n'.join(code_lines))
                    p.style = 'Normal'
                    run = p.runs[0]
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        # Skip empty lines (but add spacing)
        if not line.strip():
            doc.add_paragraph()
            continue
        
        # Handle headings
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        
        # Handle horizontal rules
        elif line.strip() == '---':
            p = doc.add_paragraph('_' * 80)
            run = p.runs[0]
            run.font.size = Pt(10)
        
        # Handle bold **text**
        elif '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        
        # Handle lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Remove markdown bold from list items
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
        
        elif re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
        
        # Handle tables (basic support)
        elif line.strip().startswith('|'):
            # Skip markdown table formatting for now
            if '---' not in line:
                p = doc.add_paragraph(line.strip())
                run = p.runs[0]
                run.font.size = Pt(10)
        
        # Regular paragraphs
        else:
            # Remove markdown bold syntax
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            # Remove inline code backticks
            text = re.sub(r'`(.*?)`', r'\1', text)
            p = doc.add_paragraph(text)
            run = p.runs[0] if p.runs else None
            if run:
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
    
    # Save document
    doc.save(docx_file)
    
    # Get page count estimation
    # Word: ~250-300 words per page with 11pt Times New Roman, 1" margins
    total_chars = sum(len(line) for line in lines)
    estimated_words = total_chars / 6  # rough estimate
    estimated_pages = estimated_words / 250  # conservative estimate
    
    print(f"\n✅ Conversion complete!")
    print(f"📄 Output: {docx_file}")
    print(f"📊 Stats:")
    print(f"   - Lines: {len(lines)}")
    print(f"   - Characters: {total_chars:,}")
    print(f"   - Estimated words: {int(estimated_words):,}")
    print(f"   - Estimated pages (250 words/page): ~{int(estimated_pages)} pages")
    print(f"\n💡 Open in Word to see actual page count with formatting")

if __name__ == '__main__':
    import sys
    
    md_file = r'f:\FINALPROJECT\Codemaster-3\docs\academicreport\version2.md'
    docx_file = r'f:\FINALPROJECT\Codemaster-3\docs\academicreport\version2.docx'
    
    try:
        convert_md_to_docx(md_file, docx_file)
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
