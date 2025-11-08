import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def convert_md_to_docx(md_file, docx_file):
    """Convert Markdown file to DOCX with formatting"""
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Process content line by line
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('##### '):
            doc.add_heading(line[6:], level=5)
        
        # Horizontal rule
        elif line.startswith('---') or line.startswith('***'):
            doc.add_paragraph('_' * 50)
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
        
        # Code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text)
            p.style = 'No Spacing'
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        # Bold text
        elif '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        
        # Regular paragraph
        else:
            doc.add_paragraph(line)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✅ Đã chuyển đổi thành công: {docx_file}")

if __name__ == '__main__':
    md_file = 'version3.md'
    docx_file = 'version3.docx'
    
    try:
        convert_md_to_docx(md_file, docx_file)
    except Exception as e:
        print(f"❌ Lỗi: {e}")
